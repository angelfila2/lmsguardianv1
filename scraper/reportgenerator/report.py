import os
import requests
import smtplib
from datetime import datetime, UTC
from typing import List
from docx import Document
from email.message import EmailMessage
from dotenv import load_dotenv
import pytz

load_dotenv()


def generatePDF(ucname: str, moduleCode: str, urls: List[dict],baseUrl:str) -> str:
    # Load Word template
    template_path = r"C:\Users\Asus\OneDrive - Murdoch University\Desktop\LMSGuardian\scraper\reportgenerator\templateReportUC.docx"
    doc = Document(template_path)

    # Replace placeholders
    replacements = {
        "<Name>": ucname,
        "<Module code>": moduleCode,
        "<URL>": baseUrl,
    }

    for para in doc.paragraphs:
        for key, value in replacements.items():
            if key in para.text:
                para.text = para.text.replace(key, value)

    # Fill the first table in the template
    tables = doc.tables
    if tables:
        table = tables[0]
        for link in urls:
            row = table.add_row().cells
            row[0].text = link.get("url_link", "")
            row[1].text = link.get("risk_status", "")
            row[2].text = link.get("scraped_at", "")

    # Save with timestamped filename
    sg = pytz.timezone("Asia/Singapore")
    safe_code = datetime.now(sg).strftime("%Y-%m-%d")
    print(safe_code + "_"+moduleCode)
    filename = f"{safe_code}_{moduleCode}_report.docx"
    output_dir = r"C:\Users\Asus\OneDrive - Murdoch University\Desktop\LMSGuardian\scraper\reportgenerator\report"
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, filename)
    doc.save(output_path)

    print(f"✅ Report saved to: {output_path}")
    return output_path


def send_email_with_report(
    to_email: str, attachment_path: str, moduleCode: str, ucname: str
):
    EMAIL_ADDRESS = os.getenv("EMAIL_USER")
    EMAIL_PASSWORD = os.getenv("EMAIL_PASS")

    if not EMAIL_ADDRESS or not EMAIL_PASSWORD:
        print("Email credentials not set in environment.")
        return

    msg = EmailMessage()
    msg["From"] = f"LMS Guardian <{EMAIL_ADDRESS}>"
    msg["To"] = to_email
    msg["Subject"] = f"[ALERT] High-risk links detected in {moduleCode}"
    msg["Reply-To"] = "noreply@example.com"

    body = f"""
Dear {ucname},

We wish to inform you that high-risk external links have been detected on the LMS course site for {moduleCode}. As the Unit Coordinator, your attention is required to review and address the issues identified in the attached report.

Please find the report enclosed for your reference.

(This is an automatically generated notification. Please do not reply.)

Best regards,  
LMS Guardian Team
"""
    msg.set_content(body)

    # Attach the .docx file
    with open(attachment_path, "rb") as f:
        msg.add_attachment(
            f.read(),
            maintype="application",
            subtype="vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename=os.path.basename(attachment_path),
        )

    try:
        with smtplib.SMTP("smtp.gmail.com", 587) as smtp:
            # smtp.set_debuglevel(
            #     1
            # )  # this gonna cause alot of output, only uncommnet if u need see wats goin on
            smtp.ehlo()
            smtp.starttls()
            smtp.ehlo()
            smtp.login(EMAIL_ADDRESS, EMAIL_PASSWORD)
            smtp.send_message(msg)

        print(f"📨 Email sent to {to_email} ✅")
    except Exception as e:
        print(f"❌ Failed to send email: {e}")


# === MAIN TEST ===
if __name__ == "__main__":
    test_email = "syafiqwork2023@gmail.com"
    test_module = "ICT302"
    test_uc = "Peter Cole"
    base_url="http://3.107.195.248/moodle/course/view.php?id=2"
    test_urls = [
        {"url_link": "https://example.com", "risk_status": "phishing", "scraped_at": "2025-06-14"},
        {"url_link": "https://another.com", "risk_status": "clean", "scraped_at": "2025-06-14"},
    ]

    report_path = generatePDF(test_uc, test_module, test_urls,base_url)
    send_email_with_report(test_email, report_path, test_module, test_uc)
